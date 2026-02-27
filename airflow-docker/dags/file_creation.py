# from airflow import DAG
# from airflow.operators.python import PythonOperator
# from datetime import datetime, timedelta

# from fpdf import FPDF
# from docx import Document
# import os

# # Default arguments for the DAG
# default_args = {
#     'owner': 'airflow',
#     'depends_on_past': False,
#     'start_date': datetime(2023, 8, 28),
#     'retries': 1,
#     'retry_delay': timedelta(minutes=5),
# }

# # Define the DAG
# dag = DAG(
#     'file_creation_tasks',
#     default_args=default_args,
#     description='A DAG to create text, Word, and PDF files',
#     schedule_interval=None,  # Manually triggered
#     catchup=False
# )

# # # Task 1: Create a text file
# # def create_text_file():
# #     with open(r"C:\Users\LENOVO\Documents\BigQuery.txt", "w") as file:
# #         file.write("This is a text file created by Airflow.")

# # # Task 2: Create a Word document
# # def create_word_file():
# #     doc = Document()
# #     doc.add_heading('Gaurav Word Document', 0)
# #     doc.add_paragraph("This is a Word document created by Airflow.")
# #     doc.save(r"C:\Users\LENOVO\Documents\gaurav.docx")

# # # Task 3: Create a PDF file
# # def create_pdf_file():
# #     pdf = FPDF()
# #     pdf.add_page()
# #     pdf.set_font("Arial", size=12)
# #     pdf.cell(200, 10, txt="Gaurav PDF Document", ln=True, align="C")
# #     pdf.cell(200, 10, txt="This is a PDF file created by Airflow.", ln=True, align="L")
# #     pdf.output(r"C:\Users\LENOVO\Documents\gaurav.pdf")
# BASE_PATH = "/opt/airflow/files"

# def read_text_file():
#     with open(f"{BASE_PATH}/input.txt", "r") as file:
#         return file.read()

# def convert_to_docx():
#     content = read_text_file()
#     doc = Document()
#     doc.add_paragraph(content)
#     doc.save(f"{BASE_PATH}/output.docx")

# def convert_to_pdf():
#     content = read_text_file()
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)

#     for line in content.split("\n"):
#         pdf.cell(200, 10, txt=line, ln=True)

#     pdf.output(f"{BASE_PATH}/output.pdf")# Define tasks

# create_text_task = PythonOperator(
#     task_id='read_text_file',
#     python_callable=read_text_file,
#     dag=dag,
# )

# create_word_task = PythonOperator(
#     task_id='convert_to_docx',
#     python_callable=convert_to_docx,
#     dag=dag,
# )

# create_pdf_task = PythonOperator(
#     task_id='convert_to_pdf',
#     python_callable=convert_to_pdf,
#     dag=dag,
# )

# # Set task dependencies
# create_text_task >> create_word_task >> create_pdf_task







# from airflow import DAG
# from airflow.operators.python import PythonOperator
# from airflow.operators.email import EmailOperator
# from datetime import datetime, timedelta

# from fpdf import FPDF
# from docx import Document
# import os

# # Base path inside Docker container
# BASE_PATH = "/opt/airflow/files"

# default_args = {
#     'owner': 'airflow',
#     'depends_on_past': False,
#     'start_date': datetime(2023, 8, 28),
#     'retries': 1,
#     'retry_delay': timedelta(minutes=5),
# }

# dag = DAG(
#     'file_creation_tasks',
#     default_args=default_args,
#     description='Convert text file to Word and PDF and send email',
#     schedule_interval=None,
#     catchup=False
# )

# # -----------------------------
# # Read Text File
# # -----------------------------
# def read_text_file(ti):
#     with open(f"{BASE_PATH}/input.txt", "r") as file:
#         content = file.read()
#     ti.xcom_push(key="file_content", value=content)


# # -----------------------------
# # Convert to DOCX
# # -----------------------------
# def convert_to_docx(ti):
#     content = ti.xcom_pull(task_ids='read_text_file', key="file_content")
#     doc = Document()
#     doc.add_paragraph(content)
#     output_path = f"{BASE_PATH}/output.docx"
#     doc.save(output_path)
#     ti.xcom_push(key="docx_path", value=output_path)


# # -----------------------------
# # Convert to PDF
# # -----------------------------
# def convert_to_pdf(ti):
#     content = ti.xcom_pull(task_ids='read_text_file', key="file_content")
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)

#     for line in content.split("\n"):
#         pdf.cell(200, 10, txt=line, ln=True)

#     output_path = f"{BASE_PATH}/output.pdf"
#     pdf.output(output_path)
#     ti.xcom_push(key="pdf_path", value=output_path)


# # -----------------------------
# # Email Task
# # -----------------------------
# send_success_email = EmailOperator(
#     task_id='send_success_email',
#     to='shivay999.ss@gmail.com',
#     subject='Airflow File Conversion Successful 🎉',
#     html_content="""
#         <h3>Hi Shivansh,</h3>
#         <p>Your File Conversion DAG has completed successfully.</p>
#         <ul>
#             <li><b>DAG:</b> {{ dag.dag_id }}</li>
#             <li><b>Run ID:</b> {{ run_id }}</li>
#             <li><b>Execution Date:</b> {{ ds }}</li>
#         </ul>
#         <p>PDF and DOCX files are attached.</p>
#         <br>
#         <p>Regards,<br>Airflow</p>
#     """,
#     files=[
#         f"{BASE_PATH}/output.docx",
#         f"{BASE_PATH}/output.pdf"
#     ],
#     dag=dag,
# )
# # -----------------------------
# # Define Tasks
# # -----------------------------
# create_text_task = PythonOperator(
#     task_id='read_text_file',
#     python_callable=read_text_file,
#     dag=dag,
# )

# create_word_task = PythonOperator(
#     task_id='convert_to_docx',
#     python_callable=convert_to_docx,
#     dag=dag,
# )

# create_pdf_task = PythonOperator(
#     task_id='convert_to_pdf',
#     python_callable=convert_to_pdf,
#     dag=dag,
# )

# # Task Flow
# create_text_task >> create_word_task >> create_pdf_task >> send_success_email




from airflow import DAG
from airflow.operators.python import PythonOperator
from airflow.operators.email import EmailOperator
from airflow.sensors.filesystem import FileSensor
from datetime import datetime, timedelta

from fpdf import FPDF
from docx import Document
import os

# Base path inside Docker container
BASE_PATH = "/opt/airflow/files"

default_args = {
    'owner': 'airflow',
    'depends_on_past': False,
    'start_date': datetime(2023, 8, 28),
    'retries': 1,
    'retry_delay': timedelta(minutes=5),
}

dag = DAG(
    'file_creation_tasks',
    default_args=default_args,
    description='Convert text file to Word and PDF and send email',
    schedule_interval=None,
    catchup=False
)

# -----------------------------
# FileSensor - Wait for input.txt
# -----------------------------
# wait_for_input_file = FileSensor(
#     task_id='wait_for_input_file',
#     filepath='/opt/airflow/files/input.txt',  # path to watch
#     fs_conn_id='fs_default',                  # filesystem connection
#     poke_interval=10,                         # check every 10 seconds
#     timeout=300,                              # fail after 5 minutes
#     mode='poke',                              # 'poke' or 'reschedule'
#     dag=dag,
# )


from airflow.sensors.python import PythonSensor
import os

def check_file_exists():
    return os.path.exists("/opt/airflow/files/input.txt")

wait_for_input_file = PythonSensor(
    task_id='wait_for_input_file',
    python_callable=check_file_exists,
    poke_interval=10,       # check every 10 seconds
    timeout=300,            # fail after 5 minutes
    mode='reschedule',      # release worker slot between checks
    dag=dag,
)



# -----------------------------
# Read Text File
# -----------------------------
def read_text_file(ti):
    with open(f"{BASE_PATH}/input.txt", "r") as file:
        content = file.read()
    ti.xcom_push(key="file_content", value=content)


# -----------------------------
# Convert to DOCX
# -----------------------------
def convert_to_docx(ti):
    content = ti.xcom_pull(task_ids='read_text_file', key="file_content")
    doc = Document()
    doc.add_paragraph(content)
    output_path = f"{BASE_PATH}/output.docx"
    doc.save(output_path)
    ti.xcom_push(key="docx_path", value=output_path)


# -----------------------------
# Convert to PDF
# -----------------------------
def convert_to_pdf(ti):
    content = ti.xcom_pull(task_ids='read_text_file', key="file_content")
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in content.split("\n"):
        pdf.cell(200, 10, txt=line, ln=True)
    output_path = f"{BASE_PATH}/output.pdf"
    pdf.output(output_path)
    ti.xcom_push(key="pdf_path", value=output_path)


# -----------------------------
# Define Tasks
# -----------------------------
create_text_task = PythonOperator(
    task_id='read_text_file',
    python_callable=read_text_file,
    dag=dag,
)

create_word_task = PythonOperator(
    task_id='convert_to_docx',
    python_callable=convert_to_docx,
    dag=dag,
)

create_pdf_task = PythonOperator(
    task_id='convert_to_pdf',
    python_callable=convert_to_pdf,
    dag=dag,
)

send_success_email = EmailOperator(
    task_id='send_success_email',
    to='shivay999.ss@gmail.com',
    subject='Airflow File Conversion Successful 🎉',
    html_content="""
        <h3>Hi Shivansh,</h3>
        <p>Your File Conversion DAG has completed successfully.</p>
        <ul>
            <li><b>DAG:</b> {{ dag.dag_id }}</li>
            <li><b>Run ID:</b> {{ run_id }}</li>
            <li><b>Execution Date:</b> {{ ds }}</li>
        </ul>
        <p>PDF and DOCX files are attached.</p>
        <br>
        <p>Regards,<br>Airflow</p>
    """,
    files=[
        f"{BASE_PATH}/output.docx",
        f"{BASE_PATH}/output.pdf"
    ],
    conn_id='smtp_default',
    dag=dag,
)

# -----------------------------
# Task Flow
# -----------------------------
wait_for_input_file >> create_text_task >> [create_word_task, create_pdf_task] >> send_success_email